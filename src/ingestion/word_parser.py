"""
word_parser.py
==============
.docx 파일 파서 (python-docx).

- 문단(paragraph) 과 표(table) 를 각각 ParsedSection 으로 반환.
- 빈 문단 제거.
- 표는 행 단위 탭 구분 텍스트로 변환.
- heading 스타일이 있으면 section_title 로 활용.
- 이미지/OCR 은 TODO.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List

from src.logger import get_logger

log = get_logger(__name__)


def _is_heading(style_name: str | None) -> bool:
    if not style_name:
        return False
    return style_name.lower().startswith("heading")


def parse_word(path: Path, document_id: str) -> List[Dict[str, Any]]:
    try:
        from docx import Document as DocxDocument  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "python-docx 가 설치되어 있지 않습니다. requirements.txt 를 다시 설치하세요."
        ) from e

    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        log.error("Word 로드 실패: %s (%s)", path.name, e)
        raise

    sections: List[Dict[str, Any]] = []
    current_heading: str | None = None
    current_paragraphs: List[str] = []
    paragraph_index = 0

    def flush_paragraphs():
        nonlocal current_paragraphs, paragraph_index
        if not current_paragraphs:
            return
        text = "\n".join(p for p in current_paragraphs if p.strip())
        if text.strip():
            sections.append({
                "section_title": current_heading,
                "content_type": "text",
                "content": text,
                "metadata": {
                    "paragraph_index": paragraph_index,
                    "paragraph_count": len(current_paragraphs),
                },
            })
            paragraph_index += 1
        current_paragraphs = []

    # 1) 문단 처리
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = getattr(p.style, "name", "") if p.style else ""
        if _is_heading(style_name):
            flush_paragraphs()
            current_heading = text
        else:
            current_paragraphs.append(text)
    flush_paragraphs()

    # 2) 표 처리
    for t_idx, table in enumerate(doc.tables):
        try:
            rows: List[str] = []
            for row in table.rows:
                cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                sections.append({
                    "section_title": current_heading or f"표 {t_idx + 1}",
                    "content_type": "table",
                    "content": "\n".join(rows),
                    "metadata": {
                        "table_index": t_idx,
                        "row_count": len(rows),
                    },
                })
        except Exception as e:  # noqa: BLE001
            log.warning("Word 표 파싱 실패: %s (table=%d, %s)", path.name, t_idx, e)
            continue

    log.info("Word 파싱 완료: %s -> %d section", path.name, len(sections))
    return sections
